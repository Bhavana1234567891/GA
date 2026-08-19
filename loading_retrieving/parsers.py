"""Turn raw file bytes into a list of {text, metadata} documents."""

from __future__ import annotations

import csv
import io
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv", ".html"}


def decode_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def parse(filename: str, data: bytes) -> list[dict]:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type '{suffix}'. Use: {supported}")

    parsers = {
        ".txt": parse_text,
        ".md": parse_text,
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".csv": parse_csv,
        ".html": parse_html,
    }
    return parsers[suffix](data, filename, suffix.lstrip("."))


def parse_text(data: bytes, source: str, file_type: str) -> list[dict]:
    text = decode_bytes(data)
    return [{"text": text, "metadata": {"source": source, "file_type": file_type}}]


def parse_pdf(data: bytes, source: str, file_type: str) -> list[dict]:
    reader = PdfReader(io.BytesIO(data))
    docs = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        docs.append(
            {
                "text": text,
                "metadata": {"source": source, "file_type": file_type, "page": index},
            }
        )
    if not docs:
        raise ValueError(f"No text extracted from '{source}'. Scanned PDFs are not supported.")
    return docs


def parse_docx(data: bytes, source: str, file_type: str) -> list[dict]:
    document = Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(cell for cell in cells if cell)
            if line:
                parts.append(line)
    return [{"text": "\n".join(parts), "metadata": {"source": source, "file_type": file_type}}]


def parse_html(data: bytes, source: str, file_type: str) -> list[dict]:
    soup = BeautifulSoup(data, "lxml")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    text = soup.get_text("\n")
    return [{"text": text, "metadata": {"source": source, "file_type": file_type}}]


def parse_csv(data: bytes, source: str, file_type: str) -> list[dict]:
    text = decode_bytes(data)
    rows = list(csv.reader(io.StringIO(text)))
    if not rows:
        return []
    header = ", ".join(cell.strip() for cell in rows[0])
    body = [", ".join(cell.strip() for cell in row) for row in rows[1:]]
    return [
        {
            "text": "\n".join(body),
            "metadata": {"source": source, "file_type": file_type, "header": header},
        }
    ]
